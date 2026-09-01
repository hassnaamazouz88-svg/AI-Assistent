from docx import Document


def extract_docx(file_path: str) -> dict:
    """
    Extrait le contenu structuré d'un fichier DOCX, sans le
    transformer en texte aplati — ce flattening est délégué à
    text_builder.py, juste avant le chunking.

    Retourne :
    {
        "paragraphes": ["texte du paragraphe 1", "texte du paragraphe 2", ...],
        "tableaux": [
            [  # tableau 1
                ["cellule1", "cellule2", ...],  # ligne 1
                ["cellule1", "cellule2", ...],  # ligne 2
            ],
            [  # tableau 2
                ...
            ]
        ],
        "nb_paragraphes": 42,
        "nb_tableaux": 2
    }
    """

    document = Document(file_path)

    # Paragraphes
    paragraphes = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphes.append(text)

    # Tableaux
    tableaux = []
    for table in document.tables:

        lignes = []
        for row in table.rows:

            cellules = [cell.text.strip() for cell in row.cells]

            if any(cellules):
                lignes.append(cellules)

        if lignes:
            tableaux.append(lignes)

    return {
        "paragraphes": paragraphes,
        "tableaux": tableaux,
        "nb_paragraphes": len(paragraphes),
        "nb_tableaux": len(tableaux)
    }